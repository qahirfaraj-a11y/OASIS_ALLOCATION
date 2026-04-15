import os
import pandas as pd
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

# Paths
SCORECARD_FILE = r"C:\Users\iLink\.gemini\antigravity\scratch\Full_Product_Allocation_Scorecard_v7.csv"
EDGES_FILE = r"C:\Users\iLink\.gemini\antigravity\scratch\neutral_network_export\edges.csv"
RATIOS_FILE = r"C:\Users\iLink\.gemini\antigravity\scratch\oasis\data\supplier_dept_ratios.json"
OUTPUT_FILE = r"C:\Users\iLink\.gemini\antigravity\scratch\OASIS_Supplier_Relational_Audit.docx"

def generate_report():
    print("Initiating Supplier Intelligence Extraction...")
    
    # 1. Load Data
    scorecard_df = pd.read_csv(SCORECARD_FILE)
    
    # Pre-process SKU-to-Supplier map
    sku_to_supp = scorecard_df.set_index('Product')['Supplier'].to_dict()
    
    # Aggregating Supplier Metrics
    supplier_stats = scorecard_df.groupby('Supplier').agg({
        'Total_Revenue': 'sum',
        'Margin_Pct': 'mean',
        'Supplier_Reliability': 'mean',
        'Capital_Required': 'sum',
        'Product': 'count',
        'Department': lambda x: sorted(list(set(x)))
    }).sort_values('Total_Revenue', ascending=False)

    # 2. Neural Affinities (Supplier-to-Supplier)
    neural_links = defaultdict(lambda: defaultdict(int))
    if os.path.exists(EDGES_FILE):
        edges_df = pd.read_csv(EDGES_FILE)
        for _, row in edges_df.iterrows():
            if row['relation'] == 'link':
                src_sku, tgt_sku = row['source'], row['target']
                src_supp = sku_to_supp.get(src_sku)
                tgt_supp = sku_to_supp.get(tgt_sku)
                if src_supp and tgt_supp and src_supp != tgt_supp:
                    neural_links[src_supp][tgt_supp] += 1
                    neural_links[tgt_supp][src_supp] += 1

    # 3. Market Rivals (Category Overlap)
    market_rivals = defaultdict(set)
    if os.path.exists(RATIOS_FILE):
        with open(RATIOS_FILE, 'r') as f:
            dept_map = json.load(f)
        for dept, supps in dept_map.items():
            s_list = [s for s in supps.keys() if s in supplier_stats.index]
            for s1 in s_list:
                for s2 in s_list:
                    if s1 != s2:
                        market_rivals[s1].add(s2)

    # 4. Generate Word Document
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title Page
    title = doc.add_heading('O.A.S.I.S. Supplier Relational Audit', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Confidential Intelligence Report - Deep Network Connectivity Analysis').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 5)
    
    # Executive Summary
    doc.add_heading('Executive Network Summary', level=1)
    summary = (
        f"This audit encompasses {len(supplier_stats)} active suppliers managing a total network portfolio "
        f"of {int(supplier_stats['Product'].sum())} SKUs. The O.A.S.I.S. neural network has identified "
        f"{sum(len(v) for v in neural_links.values()) // 2} distinct inter-supplier affinities based on "
        f"product-level demand correlations."
    )
    doc.add_paragraph(summary)
    doc.add_page_break()

    # 5. Continuous Dossier Generation
    print(f"Generating Dossiers for {len(supplier_stats)} suppliers...")
    
    for supp_name, row in supplier_stats.iterrows():
        # Supplier Header
        h = doc.add_heading(str(supp_name), level=2)
        
        # Financial Overview Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Metric'
        hdr_cells[3].text = 'Value'
        
        metrics = [
            ('Dominance (Revenue)', f"KES {row['Total_Revenue']:,.2f}"),
            ('Reliability Index', f"{row['Supplier_Reliability']*100:.1f}%"),
            ('SKU Count', str(int(row['Product']))),
            ('Portfolio Margin', f"{row['Margin_Pct']:.2f}%"),
            ('Capital Commitment', f"KES {row['Capital_Required']:,.2f}"),
            ('Primary Categories', ", ".join(row['Department'][:3]))
        ]
        
        for i in range(0, len(metrics), 2):
            row_cells = table.add_row().cells
            row_cells[0].text = metrics[i][0]
            row_cells[1].text = metrics[i][1]
            if i+1 < len(metrics):
                row_cells[2].text = metrics[i+1][0]
                row_cells[3].text = metrics[i+1][1]

        # Neural Affinities
        doc.add_heading('Neural Network Affinity', level=3)
        affs = sorted(neural_links[supp_name].items(), key=lambda x: x[1], reverse=True)[:5]
        if affs:
            p = doc.add_paragraph()
            p.add_run("Discovered high-affinity partners (Neural Network Correlation):").bold = True
            for partner, weight in affs:
                doc.add_paragraph(f"• {partner} (Connectivity Weight: {weight})", style='List Bullet')
        else:
            doc.add_paragraph("No significant neural affinities discovered in the current training set.")

        # Market Rivals
        doc.add_heading('Market Rivalry (Direct Competitors)', level=3)
        rivals = list(market_rivals[supp_name])[:8]
        if rivals:
            p = doc.add_paragraph()
            p.add_run("Primary category competitors sharing shelf-space dominance:").bold = True
            rival_text = ", ".join(rivals)
            doc.add_paragraph(rival_text)
        else:
            doc.add_paragraph("No direct category rivals identified (Niche/Exclusive Provider).")

        doc.add_paragraph("_" * 50) # Spacer

    # Save
    doc.save(OUTPUT_FILE)
    print(f"Report successfully saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    generate_report()
