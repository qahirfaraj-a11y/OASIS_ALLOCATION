import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_base_doc(title_text, subtitle_text):
    doc = Document()
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(subtitle_text, 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def generate_unique_moats():
    print("Generating Unique Moats Document...")
    df = pd.read_csv("unique_skus_analysis.csv").sort_values(by='moat_score', ascending=False)
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "The Unique Moat Portfolio (Zero Substitutes)")
    doc.add_paragraph("Focus: Identifying SKUs with absolute monopolistic demand. Stockouts here cause 100% revenue loss.")
    
    for _, row in df.head(50).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"UNIQUE MOAT: {row['id']}")
        run.bold = True
        doc.add_paragraph(
            f"Role: Monopolistic Anchor in the {row['department']} category. "
            f"Strategy: Maintain 100% Fill Rate. No alternative brands exist for this SKU in the neural graph."
        )
    doc.save("OASIS_Strategy_Unique_Moats.docx")

def generate_backstop_anchors():
    print("Generating Backstop Anchors Document...")
    df = pd.read_csv("backstop_anchors_analysis.csv").sort_values(by='backstop_count', ascending=False)
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "The Backstop Anchor Portfolio (Universal Substitutes)")
    doc.add_paragraph("Focus: SKUs that act as hubs, recapturing revenue from hundreds of other category stockouts.")
    
    for _, row in df.head(50).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"BACKSTOP ANCHOR: {row['id']}")
        run.bold = True
        doc.add_paragraph(
            f"Hub Potential: This SKU is a valid substitute for {row['backstop_count']} other items. "
            f"Strategy: Prioritize as the 'Backstop' to safeguard category-wide revenue stability."
        )
    doc.save("OASIS_Strategy_Backstop_Anchors.docx")

def generate_supplier_fragility():
    print("Generating Supplier Fragility Document...")
    df = pd.read_csv("supplier_fragility_map.csv").sort_values(by='ads', ascending=False)
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "Supplier Fragility & Risk Mitigation Map")
    doc.add_paragraph("Focus: Identifying high-volume clusters where the entire substitution network is controlled by one supplier.")
    
    for _, row in df.head(50).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"FRAGILITY ALERT: {row['id']}")
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        doc.add_paragraph(
            f"Vulnerability: {row['monopoly_supplier']} controls 100% of the {row['substitute_count']} available substitutes. "
            f"Strategy: Introduce 'Monopoly Breaker' imports to diversify the fallback network."
        )
    doc.save("OASIS_Strategy_Supplier_Fragility.docx")

def generate_private_label():
    print("Generating Private Label Opportunity Document...")
    df = pd.read_csv("private_label_index.csv").sort_values(by='substitution_density', ascending=False)
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "Private Label Opportunity & Commoditization Index")
    doc.add_paragraph("Focus: Categories with extremely high substitution density, indicating low brand loyalty.")
    
    for _, row in df.head(50).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"PL TARGET: {row['department']}")
        run.bold = True
        doc.add_paragraph(
            f"Commoditization Factor: Substitution density is {row['substitution_density']:.1f}/5.0 (Extreme Interchangeability). "
            f"Strategy: Ideal category for High-Margin Private Label import displacement."
        )
    doc.save("OASIS_Strategy_Private_Label_Index.docx")

if __name__ == "__main__":
    generate_unique_moats()
    generate_backstop_anchors()
    generate_supplier_fragility()
    generate_private_label()
    print("All isolated strategic documents created successfully.")
