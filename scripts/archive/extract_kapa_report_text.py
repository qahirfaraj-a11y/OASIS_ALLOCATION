import docx

doc = docx.Document(r"C:\Users\iLink\.gemini\antigravity\scratch\Kapa_Oil_Assessment_Report.docx")
text = []
for p in doc.paragraphs:
    text.append(p.text)

print("=== PARAGRAPHS ===")
for line in text:
    if any(keyword in line.upper() for keyword in ["TRANSPORT", "LOGISTICS", "FREIGHT", "LEAD TIME", "DELIVERY", "ROUTE", "FLEET"]):
        print(line)

print("=== TABLES ===")
for t in doc.tables:
    for r in t.rows:
        row_text = [c.text.strip() for c in r.cells]
        row_str = " | ".join(row_text)
        if "KAPA" in row_str.upper() or any(keyword in row_str.upper() for keyword in ["TRANSPORT", "LOGISTICS", "FREIGHT", "LEAD TIME", "DELIVERY"]):
            print(row_text)
