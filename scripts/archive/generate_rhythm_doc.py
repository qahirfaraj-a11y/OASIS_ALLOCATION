import json
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_supplier_rhythm_doc():
    json_path = r'C:\Users\iLink\.gemini\antigravity\scratch\supplier_rhythm_analysis.json'
    output_path = r'C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Rhythm_Master.docx'

    if not os.path.exists(json_path):
        print(f"Error: {json_path} not found.")
        return

    with open(json_path, 'r') as f:
        data = json.load(f)

    rhythm_data = data.get('po_rhythm', {})

    doc = Document()

    # Title
    title = doc.add_heading('O.A.S.I.S. Supplier Rhythm Master', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This document contains the calculated delivery and ordering rhythms for all registered suppliers, used by the O.A.S.I.S. forecasting engines to optimize procurement timing.')

    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Supplier Name'
    hdr_cells[1].text = 'Median Gap (Days)'
    hdr_cells[2].text = 'Avg Gap (Days)'
    hdr_cells[3].text = 'Total Orders'
    hdr_cells[4].text = 'Last Order Date'

    # Set bold for header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Sorting suppliers by name
    sorted_suppliers = sorted(rhythm_data.keys())

    for supplier in sorted_suppliers:
        s_data = rhythm_data[supplier]
        row_cells = table.add_row().cells
        row_cells[0].text = str(supplier)
        row_cells[1].text = str(s_data.get('median_gap', 'N/A'))
        row_cells[2].text = f"{s_data.get('avg_gap', 0):.2f}" if isinstance(s_data.get('avg_gap'), (int, float)) else str(s_data.get('avg_gap'))
        row_cells[3].text = str(s_data.get('total_orders', '0'))
        row_cells[4].text = str(s_data.get('last_order', 'N/A'))

    # Save
    doc.save(output_path)
    print(f"Success: Document saved to {output_path}")

if __name__ == "__main__":
    generate_supplier_rhythm_doc()
