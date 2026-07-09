import docx
import os

doc_path = r"C:\Users\iLink\.gemini\antigravity\scratch\Kapa_Portfolio_Specific_Deep_Dive_v2.docx"
if os.path.exists(doc_path):
    print("Document found!")
    doc = docx.Document(doc_path)
    print(f"Total tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        # Print first 5 rows
        for r_idx, row in enumerate(table.rows[:5]):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            print(f"  Row {r_idx}: {cells}")
else:
    print("Document not found at:", doc_path)
