import docx

def dump_docx(f):
    doc = docx.Document(f)
    print(f"=== {f} ===")
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
    for t in doc.tables:
        for r in t.rows:
            print([c.text.strip() for c in r.cells])

print("Dumping Kapa_Oil_Assessment_Report.docx:")
try:
    dump_docx(r"C:\Users\iLink\.gemini\antigravity\scratch\Kapa_Oil_Assessment_Report.docx")
except Exception as e:
    print(e)

print("\nDumping Unique_SKU_Ecosystem_Analysis.docx:")
try:
    dump_docx(r"C:\Users\iLink\.gemini\antigravity\scratch\Unique_SKU_Ecosystem_Analysis.docx")
except Exception as e:
    print(e)
