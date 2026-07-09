import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_docs():
    csv_input = "unique_skus_analysis.csv"
    excel_output = "unique_moats_portfolio.xlsx"
    word_output = "Unique_SKU_Ecosystem_Analysis.docx"
    
    print(f"Loading data from {csv_input}...")
    df = pd.read_csv(csv_input)
    
    # --- 1. Excel Export ---
    print("Generating Excel Portfolio...")
    with pd.ExcelWriter(excel_output, engine='openpyxl') as writer:
        # Full Data
        df.to_excel(writer, sheet_name='Unique SKUs', index=False)
        
        # Departmental Summary
        dept_summary = df.groupby('department').agg({
            'id': 'count',
            'moat_score': 'sum',
            'velocity_ads': 'mean'
        }).rename(columns={'id': 'sku_count'}).sort_values(by='moat_score', ascending=False)
        dept_summary.to_excel(writer, sheet_name='Category Moats')
        
    print(f"Excel file created: {excel_output}")
    
    # --- 2. Word Document Export ---
    print("Generating Word Strategic Guide...")
    doc = Document()
    
    # Title
    title = doc.add_heading('O.A.S.I.S. Strategic Guide:', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('The Unique SKU Ecosystem (Monopolistic Moats)', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This document analyzes SKUs identified by the O.A.S.I.S. Neural Network as having zero recorded substitutes. "
        "Each item listed here represents a 'Unique Moat'—an island of demand that captures store traffic that cannot be "
        "diverted to other products in the event of a stockout."
    )
    
    # Group by Department for the Word Doc
    departments = sorted(df['department'].unique())
    
    for dept in departments:
        doc.add_heading(f'Category: {dept.upper()}', level=1)
        dept_items = df[df['department'] == dept].sort_values(by='moat_score', ascending=False)
        
        for index, row in dept_items.iterrows():
            # SKU Header
            p = doc.add_paragraph()
            run = p.add_run(f"SKU: {row['id']}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Key Stats Table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = f"Supplier: {row['supplier']}"
            hdr_cells[1].text = f"Price: {row['price']:.2f}"
            hdr_cells[2].text = f"ADS: {row['velocity_ads']:.4f}"
            hdr_cells[3].text = f"Moat Score: {row['moat_score']:.2f}"
            
            # Narrative Generation
            v = row['velocity_ads']
            if v > 1.0:
                status = "Dominant Anchor"
                reasoning = "a high-velocity essential driver of category traffic."
            elif v > 0.1:
                status = "Stable Moat"
                reasoning = "a consistent volume contributor with zero cross-brand competition."
            else:
                status = "Niche Specialist"
                reasoning = "a highly specialized SKU capturing a non-substitutable customer segment."
            
            narrative = (
                f"Identified as a {status} in the {dept} ecosystem from our neural network. "
                f"With a daily velocity of {row['velocity_ads']:.2f}, this SKU represents {reasoning} "
                f"Because our neural network recognizes zero substitutes for this SKU, maintaining its availability "
                f"is critical for protecting your gross profit of KES {row['gross_profit']:.2f}."
            )
            
            doc.add_paragraph(narrative)
            doc.add_paragraph("") # Spacer
            
    doc.save(word_output)
    print(f"Word document created: {word_output}")

if __name__ == "__main__":
    generate_docs()
