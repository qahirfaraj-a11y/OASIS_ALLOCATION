import pandas as pd
import json
import os
import docx

def analyze_kapa_data():
    kapa_term = "KAPA"
    results = {}

    # 1. kapa.xlsx
    kapa_xlsx_path = r"C:\Users\iLink\Downloads\kapa.xlsx"
    try:
        df_kapa = pd.read_excel(kapa_xlsx_path)
        print(f"--- {kapa_xlsx_path} ---")
        print(df_kapa.head())
        print(f"Total Rows: {len(df_kapa)}\n")
    except Exception as e:
        print(f"Error reading {kapa_xlsx_path}: {e}")

    # 2. neutral_network_export/nodes.csv
    nodes_csv = r"C:\Users\iLink\.gemini\antigravity\scratch\neutral_network_export\nodes.csv"
    try:
        df_nodes = pd.read_csv(nodes_csv)
        kapa_nodes = df_nodes[df_nodes.astype(str).apply(lambda x: x.str.contains(kapa_term, case=False, na=False)).any(axis=1)]
        print(f"--- {nodes_csv} ---")
        print(f"Found {len(kapa_nodes)} KAPA nodes")
        if not kapa_nodes.empty:
            print(kapa_nodes.head())
        print()
    except Exception as e:
        print(f"Error reading {nodes_csv}: {e}")

    # 3. neutral_network_export/edges.csv
    edges_csv = r"C:\Users\iLink\.gemini\antigravity\scratch\neutral_network_export\edges.csv"
    try:
        df_edges = pd.read_csv(edges_csv)
        kapa_edges = df_edges[df_edges.astype(str).apply(lambda x: x.str.contains(kapa_term, case=False, na=False)).any(axis=1)]
        print(f"--- {edges_csv} ---")
        print(f"Found {len(kapa_edges)} KAPA edges")
        print()
    except Exception as e:
        print(f"Error reading {edges_csv}: {e}")

    # 4. neutral_network_export/full_graph.json
    graph_json = r"C:\Users\iLink\.gemini\antigravity\scratch\neutral_network_export\full_graph.json"
    try:
        with open(graph_json, 'r', encoding='utf-8') as f:
            graph_data = json.load(f)
        kapa_items_in_json = 0
        
        # very basic text search in json dump
        graph_str = json.dumps(graph_data).upper()
        print(f"--- {graph_json} ---")
        print(f"Graph JSON contains 'KAPA': {kapa_term in graph_str}")
        print()
    except Exception as e:
        print(f"Error reading {graph_json}: {e}")

    # 5. Supplier_Fulfillment_Summary.xlsx
    supp_fulf = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Fulfillment_Summary.xlsx"
    try:
        df_supp = pd.read_excel(supp_fulf)
        kapa_supp = df_supp[df_supp.astype(str).apply(lambda x: x.str.contains(kapa_term, case=False, na=False)).any(axis=1)]
        print(f"--- {supp_fulf} ---")
        print(f"Found {len(kapa_supp)} KAPA rows")
        if not kapa_supp.empty:
            print(kapa_supp)
        print()
    except Exception as e:
        print(f"Error reading {supp_fulf}: {e}")

    # 6. Supplier_Intelligence_Report_2025_v3.xlsx
    supp_intel = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Intelligence_Report_2025_v3.xlsx"
    try:
        df_intel = pd.read_excel(supp_intel)
        kapa_intel = df_intel[df_intel.astype(str).apply(lambda x: x.str.contains(kapa_term, case=False, na=False)).any(axis=1)]
        print(f"--- {supp_intel} ---")
        print(f"Found {len(kapa_intel)} KAPA rows")
        if not kapa_intel.empty:
            print(kapa_intel)
        print()
    except Exception as e:
        print(f"Error reading {supp_intel}: {e}")

    # 7. Supplier_Master_Intelligence_Report.docx
    supp_master_docx = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx"
    try:
        doc = docx.Document(supp_master_docx)
        print(f"--- {supp_master_docx} ---")
        kapa_paragraphs = [p.text for p in doc.paragraphs if kapa_term in p.text.upper()]
        print(f"Found {len(kapa_paragraphs)} paragraphs mentioning KAPA.")
        for p in kapa_paragraphs[:5]:
            print(f"- {p[:150]}...")
        print()
    except Exception as e:
        print(f"Error reading {supp_master_docx}: {e}")

    # 8. Supplier_Rhythm_Master.docx
    supp_rhythm_docx = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Rhythm_Master.docx"
    try:
        doc = docx.Document(supp_rhythm_docx)
        print(f"--- {supp_rhythm_docx} ---")
        kapa_paragraphs = [p.text for p in doc.paragraphs if kapa_term in p.text.upper()]
        print(f"Found {len(kapa_paragraphs)} paragraphs mentioning KAPA.")
        for p in kapa_paragraphs[:5]:
            print(f"- {p[:150]}...")
        print()
    except Exception as e:
        print(f"Error reading {supp_rhythm_docx}: {e}")

if __name__ == "__main__":
    analyze_kapa_data()
