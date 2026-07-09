import docx
import os

def search():
    f = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Rhythm_Master.docx"
    if not os.path.exists(f):
        print("Supplier_Rhythm_Master.docx not found")
        return
    doc = docx.Document(f)
    print("=== Supplier_Rhythm_Master.docx transport info ===")
    for i, p in enumerate(doc.paragraphs):
        if "TRANSPORT" in p.text.upper() or "LOGISTICS" in p.text.upper() or "FREIGHT" in p.text.upper():
            print(f"P{i}: {p.text}")
    for t_idx, t in enumerate(doc.tables):
        for r_idx, r in enumerate(t.rows):
            row_text = [c.text.strip() for c in r.cells]
            row_str = " | ".join(row_text)
            if "TRANSPORT" in row_str.upper() or "LOGISTICS" in row_str.upper() or "FREIGHT" in row_str.upper():
                print(f"Table {t_idx} Row {r_idx}: {row_text}")

search()
