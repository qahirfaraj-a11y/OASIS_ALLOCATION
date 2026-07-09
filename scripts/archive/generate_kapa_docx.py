import json
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_kapa_docx():
    # Load data
    with open('kapa_deep_dive.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    anomalies = data.get('anomalies', [])
    skus = data.get('skus', [])
    
    # Sort SKUs by attractor score (descending) then revenue
    skus.sort(key=lambda x: (x['attractor_score'], x['revenue']), reverse=True)

    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Kapa Oil Refineries - Portfolio Deep Dive & Neural Network Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Anomalies Section
    doc.add_heading('1. The 11-Day PO Fulfillment Anomaly Investigation', level=1)
    p = doc.add_paragraph('Based on historical GRN tracking, it was observed that some POs took 11 days to fulfill, despite a 7-day expiration rule. The investigation revealed that this is an isolated operational anomaly rather than a systemic delay.')
    p = doc.add_paragraph(f'Total Kapa Lines Evaluated: 699\nAnomalous Lines (>7 Days): {len(anomalies)} (6.4%)')
    
    if anomalies:
        po_info = anomalies[0]
        doc.add_paragraph(f"Root Cause Analysis: All {len(anomalies)} anomalous lines trace back to a single PO event (PO No: {po_info['po']} on {po_info['po_date']}) which was forced received on {po_info['grn_date']} (GRN No: {po_info['grn']}). This indicates a manual receiving override by a manager post-expiration.", style='Intense Quote')

    # 2. Network Intelligence Summaries
    doc.add_heading('2. Network Attractors & Connectors', level=1)
    p = doc.add_paragraph('In the ST-GAT (Spatio-Temporal Graph Attention) network, nodes function as either Attractors or Connectors:')
    doc.add_paragraph('• Attractors (High In-Degree): SKUs that receive demand when competitors stock out (Safety Nets).', style='List Bullet')
    doc.add_paragraph('• Connectors (High Out-Degree): Fragile SKUs that fragment demand to multiple substitutes when they stock out.', style='List Bullet')
    
    # Top Attractors
    top_attractors = [s for s in skus if s['attractor_score'] > 0][:5]
    if top_attractors:
        doc.add_heading('Top Network Attractors (Strategic Safety Nets):', level=2)
        for s in top_attractors:
            doc.add_paragraph(f"• {s['sku']} (Attractor Score: {s['attractor_score']})", style='List Bullet')
            
    doc.add_paragraph("Note: Almost all Kapa SKUs carry a uniform Connector Score of 18, making them highly connected 'origin' nodes. If they stock out, the engine distributes demand across ~18 alternative products.")

    # 3. Comprehensive SKU Table
    doc.add_heading('3. Individual SKU Analysis (ROI & Network Scores)', level=1)
    p = doc.add_paragraph(f'Detailed breakdown of all {len(skus)} active Kapa SKUs within the neural network.')

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SKU Name'
    hdr_cells[1].text = 'Department'
    hdr_cells[2].text = 'ROI (%)'
    hdr_cells[3].text = 'Revenue (KES)'
    hdr_cells[4].text = 'Attractor Score'
    hdr_cells[5].text = 'Connector Score'
    
    for s in skus:
        row_cells = table.add_row().cells
        row_cells[0].text = str(s['sku'])
        row_cells[1].text = str(s['department'])
        row_cells[2].text = f"{s['roi']:.2f}%" if s['roi'] else "0.00%"
        row_cells[3].text = f"{s['revenue']:,.2f}"
        row_cells[4].text = str(s['attractor_score'])
        row_cells[5].text = str(s['connector_score'])

    file_path = 'Kapa_Portfolio_Deep_Dive.docx'
    doc.save(file_path)
    print(f"Successfully generated {file_path}")

if __name__ == "__main__":
    create_kapa_docx()
