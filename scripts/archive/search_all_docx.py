import docx
import os
import glob

def search_all_docx():
    for f in glob.glob(r"C:\Users\iLink\.gemini\antigravity\scratch\*.docx"):
        if "~$" in f:
            continue
        try:
            doc = docx.Document(f)
            text = []
            for p in doc.paragraphs:
                text.append(p.text)
            for t in doc.tables:
                for r in t.rows:
                    text.append(" | ".join([c.text.strip() for c in r.cells]))
            full_text = "\n".join(text)
            if "KAPA" in full_text.upper():
                print(f"=== KAPA found in {os.path.basename(f)} ===")
                for line in text:
                    if "KAPA" in line.upper() and any(k in line.upper() for k in ["TRANSPORT", "LOGISTICS", "FREIGHT", "DELIVERY", "COST", "RATE", "LEAD"]):
                        print(f"  Line: {line}")
        except Exception as e:
            print(f"Error {f}: {e}")

search_all_docx()
