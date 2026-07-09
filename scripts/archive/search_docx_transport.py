import docx

doc = docx.Document(r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx")
print("=== Searching transport keywords in docx paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if any(k in txt.upper() for k in ["TRANSPORT", "LOGISTICS", "FREIGHT", "COST", "FEE", "RATE"]):
        print(f"P{i}: {txt}")

for t_idx, t in enumerate(doc.tables):
    for r_idx, r in enumerate(t.rows):
        row_text = [c.text.strip() for c in r.cells]
        row_str = " | ".join(row_text)
        if any(k in row_str.upper() for k in ["TRANSPORT", "LOGISTICS", "FREIGHT", "FEE", "RATE"]):
            print(f"Table {t_idx} Row {r_idx}: {row_text}")
