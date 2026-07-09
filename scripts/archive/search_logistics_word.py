import docx
import os

def search():
    f = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx"
    if not os.path.exists(f):
        print("Supplier_Master_Intelligence_Report.docx not found")
        return
    doc = docx.Document(f)
    print("=== Supplier_Master_Intelligence_Report.docx transport info ===")
    for i, p in enumerate(doc.paragraphs):
        if "TRANSPORT" in p.text.upper() or "LOGISTICS" in p.text.upper() or "FREIGHT" in p.text.upper() or "SHIPPING" in p.text.upper():
            print(f"P{i}: {p.text}")
            # print surrounding paragraphs
            for j in range(max(0, i-1), min(len(doc.paragraphs), i+2)):
                print(f"  P{j}: {doc.paragraphs[j].text}")

search()
