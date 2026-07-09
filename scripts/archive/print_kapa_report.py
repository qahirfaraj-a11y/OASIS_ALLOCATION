import docx

doc = docx.Document(r"C:\Users\iLink\.gemini\antigravity\scratch\Kapa_Oil_Assessment_Report.docx")
print("=== Kapa_Oil_Assessment_Report.docx All Content ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i}: {p.text}")
for t_idx, t in enumerate(doc.tables):
    print(f"--- Table {t_idx} ---")
    for r in t.rows:
        print([c.text.strip() for c in r.cells])
