import docx
import os

def check_doc(fn):
    f = os.path.join(r"C:\Users\iLink\.gemini\antigravity\scratch", fn)
    doc = docx.Document(f)
    print(f"\n================ {fn} ================")
    for i, p in enumerate(doc.paragraphs):
        if "KAPA" in p.text.upper():
            print(f"--- MATCH P{i} ---")
            for j in range(max(0, i-2), min(len(doc.paragraphs), i+15)):
                print(f"P{j}: {doc.paragraphs[j].text}")
    for t_idx, t in enumerate(doc.tables):
        for r_idx, r in enumerate(t.rows):
            row_text = [c.text.strip() for c in r.cells]
            row_str = " | ".join(row_text)
            if "KAPA" in row_str.upper():
                print(f"Table {t_idx} Row {r_idx}: {row_text}")

check_doc("Supplier_Rhythm_Master.docx")
check_doc("Supplier_Master_Intelligence_Report.docx")
check_doc("OASIS_Supplier_Relational_Audit.docx")
