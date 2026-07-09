import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_advanced_docs():
    backstop_csv = "backstop_anchors_analysis.csv"
    fragility_csv = "supplier_fragility_map.csv"
    commod_csv = "private_label_index.csv"
    
    excel_output = "Advanced_Import_Strategies.xlsx"
    word_output = "OASIS_Advanced_Assortment_Guide.docx"
    
    print("Loading strategy data...")
    backstop_df = pd.read_csv(backstop_csv)
    fragility_df = pd.read_csv(fragility_csv)
    commod_df = pd.read_csv(commod_csv)
    
    # --- 1. Excel Export ---
    print("Generating Advanced Excel Toolkit...")
    with pd.ExcelWriter(excel_output, engine='openpyxl') as writer:
        backstop_df.to_excel(writer, sheet_name='Backstop Anchors', index=False)
        fragility_df.to_excel(writer, sheet_name='Supplier Fragility', index=False)
        commod_df.to_excel(writer, sheet_name='Private Label Index', index=False)
    print(f"Excel file created: {excel_output}")
    
    # --- 2. Word Document Export ---
    print("Generating Advanced Strategic Guide...")
    doc = Document()
    
    # Title
    title = doc.add_heading('O.A.S.I.S. Topographical Strategy Interface:', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Advanced Import & Portfolio Allocation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This guide translates complex neural network topographical data (Centrality, Entropy, and Density) "
        "into actionable import decisions. We move beyond simple inventory to 'Network Resilience'."
    )
    
    # --- SECTION 1: BACKSTOP ANCHORS ---
    doc.add_heading('1. Revenue Protection: The Backstop Anchors', level=1)
    doc.add_paragraph(
        "These items are 'Hubs' in the substitution network. Each acts as a safety net for a massive number of other items. "
        "Protecting these SKUs is the single most efficient way to maintain 95% category revenue during broad supplier stockouts."
    )
    
    for _, row in backstop_df.head(10).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"ANCHOR: {row['id']}")
        run.bold = True
        
        narrative = (
            f"This SKU is a universal backstop for {int(row['backstop_count'])} other items in the {row['department']} category. "
            f"Strategy: Maintain 100% service level. A stockout here causes a cascade of lost revenue across the entire categories it supports."
        )
        doc.add_paragraph(narrative)

    # --- SECTION 2: SUPPLIER VULNERABILITY ---
    doc.add_heading('2. Risk Mitigation: Supplier Fragility Map', level=1)
    doc.add_paragraph(
        "This map identifies high-velocity SKU clusters where the substation network is trapped by a single supplier. "
        "Your fallback logic for these items is geographically and operationally fragile."
    )
    
    for _, row in fragility_df.head(10).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"VULNERABILITY: {row['id']}")
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0) # Red for risk
        
        narrative = (
            f"Vulnerability Detected: {row['monopoly_supplier']} controls 100% of the {int(row['substitute_count'])} substitutes "
            f"recorded for this high-volume SKU (ADS: {row['ads']:.2f}). "
            f"Strategy: Target this cluster for independent imports to break supplier monopoly and improve margins."
        )
        doc.add_paragraph(narrative)

    # --- SECTION 3: PRIVATE LABEL OPPORTUNITY ---
    doc.add_heading('3. Margin Expansion: Private Label Opportunity Index', level=1)
    doc.add_paragraph(
        "Based on substitution density. Categories with high density indicate that customers are psychologically "
        "primed to swap brands easily. These are the optimal targets for private label introductions."
    )
    
    for _, row in commod_df.head(10).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"TARGET CATEGORY: {row['department']}")
        run.bold = True
        
        narrative = (
            f"The substitution density in this category is {row['substitution_density']:.1f}/5.0. "
            f"With {int(row['sku_count'])} interchangeable SKUs and a total category velocity of {row['total_velocity']:.2f}, "
            f"this is a high-confidence zone for launching a high-margin Private Label import."
        )
        doc.add_paragraph(narrative)
        
    doc.save(word_output)
    print(f"Word document created: {word_output}")

if __name__ == "__main__":
    generate_advanced_docs()
