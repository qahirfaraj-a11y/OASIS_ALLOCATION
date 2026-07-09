"""
Convert OASIS_Exchange_Whitepaper_v1.md → .docx
Parses markdown structure and renders into a professionally styled Word document.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT  = r"C:\Users\iLink\Documents\OASIS_Exchange_Whitepaper_v1.md"
OUTPUT = r"C:\Users\iLink\Documents\OASIS_Exchange_Whitepaper_v1.docx"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_table_header(row, bg="1B3A5C", fg="FFFFFF"):
    """Style a table header row with dark background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, bg)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(fg)
                run.font.bold = True
                run.font.size = Pt(9)


def add_styled_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_table_header(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, val in enumerate(row_data):
            if c_idx < cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                # Handle bold markers
                parts = re.split(r'\*\*(.*?)\*\*', val.strip())
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(9)
                        if j % 2 == 1:  # Bold part
                            run.bold = True
                # Alternate row shading
                if r_idx % 2 == 1:
                    set_cell_shading(cell, "F2F6FA")

    return table


def parse_table_block(lines):
    """Parse markdown table lines into header and data rows."""
    # Filter out separator lines (|---|---|)
    content_lines = [l for l in lines if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', l)]
    if not content_lines:
        return None, None

    def split_row(line):
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    header = split_row(content_lines[0])
    data = [split_row(l) for l in content_lines[1:]]
    return header, data


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, size=None):
    """Add a paragraph with inline formatting (bold markers)."""
    p = doc.add_paragraph(style=style)
    # Split on bold markers
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for j, part in enumerate(parts):
        if part:
            # Also handle backtick code
            code_parts = re.split(r'`(.*?)`', part)
            for k, cp in enumerate(code_parts):
                if cp:
                    run = p.add_run(cp)
                    if bold or j % 2 == 1:
                        run.bold = True
                    if italic:
                        run.italic = True
                    if k % 2 == 1:  # Code
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
                    elif size:
                        run.font.size = size
    return p


def main():
    with open(INPUT, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Heading styles ──
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        if level == 1:
            hs.font.size = Pt(24)
        elif level == 2:
            hs.font.size = Pt(18)
        elif level == 3:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # ── Code blocks ──
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — flush
                code_text = '\n'.join(code_lines)
                if code_lines and code_lines[0].strip().startswith(('graph ', 'stateDiagram', 'flowchart', 'sequenceDiagram')):
                    # Mermaid diagram — add as italicized note
                    p = doc.add_paragraph(style='Normal')
                    run = p.add_run('[Diagram: See rendered Markdown version for interactive Mermaid diagram]')
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    # Regular code block
                    p = doc.add_paragraph(style='Normal')
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    # Light gray background via shading
                    pPr = p._p.get_or_add_pPr()
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                    pPr.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ──
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # Flush table
            header, data = parse_table_block(table_lines)
            if header and data:
                add_styled_table(doc, header, data)
                doc.add_paragraph()  # spacing
            elif header:
                add_styled_table(doc, header, [])
                doc.add_paragraph()
            in_table = False
            table_lines = []
            # Don't increment — process current line

        # ── Blank lines ──
        if not line.strip():
            i += 1
            continue

        # ── Horizontal rules ──
        if line.strip() == '---':
            # Add a thin line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 60)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # ── Headings ──
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown link syntax from heading
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Blockquotes ──
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line.strip())
            # Handle alert syntax
            text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', r'[\1]', text)
            if text.strip():
                p = doc.add_paragraph(style='Normal')
                # Indent
                p.paragraph_format.left_indent = Cm(1)
                parts = re.split(r'\*\*(.*?)\*\*', text)
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(10)
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                        if j % 2 == 1:
                            run.bold = True
                            run.italic = False
            i += 1
            continue

        # ── Bullet points ──
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            add_formatted_paragraph(doc, text, style='List Bullet')
            i += 1
            continue

        # ── Numbered list ──
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            add_formatted_paragraph(doc, text, style='List Number')
            i += 1
            continue

        # ── Regular paragraph ──
        add_formatted_paragraph(doc, line.strip())
        i += 1

    # Flush any remaining table
    if in_table and table_lines:
        header, data = parse_table_block(table_lines)
        if header and data:
            add_styled_table(doc, header, data)

    # ── Add footer ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nVersion 1.0 — Draft for Consultation — 15 May 2026')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT)
    print(f"✅ Saved: {OUTPUT}")


if __name__ == '__main__':
    main()
