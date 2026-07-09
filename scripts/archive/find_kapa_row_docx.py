import docx
import os

def check():
    f = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx"
    doc = docx.Document(f)
    print("=== Searching Kapa in Supplier_Master_Intelligence_Report.docx ===")
    
    # Check paragraphs
    for i, p in enumerate(doc.paragraphs):
        if "KAPA" in p.text.upper():
            print(f"Paragraph {i}: {p.text}")
            
    # Check tables
    for t_idx, t in enumerate(doc.tables):
        for r_idx, r in enumerate(t.rows):
            row_text = [c.text.strip() for c in r.cells]
            row_str = " | ".join(row_text)
            if "KAPA" in row_str.upper():
                print(f"Table {t_idx} Row {r_idx}: {row_text}")

check()
