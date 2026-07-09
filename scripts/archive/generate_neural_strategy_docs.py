import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_base_doc(title_text, subtitle_text):
    doc = Document()
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading(subtitle_text, 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

def generate_substitution_rivals():
    print("Generating Substitution Rivals Guide...")
    df = pd.read_csv("neural_sku_metrics_top1000.csv")
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "Competitive Rivals: Direct Substitution Mapping")
    doc.add_paragraph("Focus: Identifying the primary 'Rival' SKUs that capture demand when a specific product is stock-out.")
    
    # Filter for items that actually HAVE substitutes
    rivals_df = df[df['Top_5_Substitutes'].notnull() & (df['Top_5_Substitutes'] != "")]
    
    for _, row in rivals_df.head(100).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"PRIMARY SKU: {row['Item_Name']}")
        run.bold = True
        doc.add_paragraph(
            f"Department: {row['Department']} | ADS: {row['ADS']:.2f}\n"
            f"Direct Competitive Rivals: {row['Top_5_Substitutes']}\n"
            f"Strategy: Ensure the 'Rival' portfolio has 100% collective availability if the primary SKU is disrupted."
        )
    doc.save("OASIS_Substitution_Rivals_Guide.docx")

def generate_neural_affinities():
    print("Generating Neural Affinities Catalog...")
    df = pd.read_csv("neural_sku_metrics_top1000.csv")
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "Neural Affinities: Network Companion Mapping")
    doc.add_paragraph("Focus: Identifying high-affinity companions for cross-selling, bundling, and cognitive shelving placement.")
    
    # Filter for items that actually HAVE affinities
    affinities_df = df[df['Top_10_Neural_Affinities'].notnull() & (df['Top_10_Neural_Affinities'] != "")]
    
    for _, row in affinities_df.head(100).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"ANCHOR SKU: {row['Item_Name']}")
        run.bold = True
        doc.add_paragraph(
            f"Department: {row['Department']} | Network companions detected in the neural graph.\n"
            f"Neural Affinities: {row['Top_10_Neural_Affinities']}\n"
            f"Strategy: Bundle these companions for high-margin promotions or place them in adjacent shelving slots."
        )
    doc.save("OASIS_Neural_Affinities_Catalog.docx")

def generate_density_index():
    print("Generating SKU Density Index Report...")
    # Rank by relation count
    df = pd.read_csv("neural_sku_metrics_top1000.csv").sort_values(by='Total_Relation_Count', ascending=False)
    doc = create_base_doc("O.A.S.I.S. Strategic Guide:", "The SKU Density Index: Critical Network Hubs")
    doc.add_paragraph("Focus: High-density items that serve as the 'connective tissue' of the store ecosystem.")
    
    for _, row in df.head(100).iterrows():
        p = doc.add_paragraph()
        run = p.add_run(f"NETWORK HUB (Rank {row.name}): {row['Item_Name']}")
        run.bold = True
        doc.add_paragraph(
            f"Centrality Score: {row['Total_Relation_Count']} Links. Department: {row['Department']}\n"
            f"Status: This SKU is highly integrated. Its movement affects the velocity of {row['Total_Relation_Count']} other products.\n"
            f"Strategy: Guard these SKUs against any price or stock volatility, as they act as cognitive anchors for entire shop-trips."
        )
    doc.save("OASIS_SKU_Density_Index_Report.docx")

if __name__ == "__main__":
    generate_substitution_rivals()
    generate_neural_affinities()
    generate_density_index()
    print("Neural network strategic documents created successfully.")
