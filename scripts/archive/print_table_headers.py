import docx

doc = docx.Document(r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx")
print("Table 0 Headers:", [c.text.strip() for c in doc.tables[0].rows[0].cells])
print("Row 247 content:", [c.text.strip() for c in doc.tables[0].rows[247].cells])
