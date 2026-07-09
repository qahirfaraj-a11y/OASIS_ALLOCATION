import docx
import os

def dump():
    f = r"C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx"
    if not os.path.exists(f):
        print("Supplier_Master_Intelligence_Report.docx not found")
        return
    doc = docx.Document(f)
    print("=== Supplier_Master_Intelligence_Report.docx Paragraphs ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P{i}: {p.text}")
    for t_idx, t in enumerate(doc.tables):
        print(f"\n--- Table {t_idx} ---")
        for r_idx, r in enumerate(t.rows):
            row_text = [c.text.strip() for c in r.cells]
            print(f"R{r_idx}: {row_text}")

dump()
