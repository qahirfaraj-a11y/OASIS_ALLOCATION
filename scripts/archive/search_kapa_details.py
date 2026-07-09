import docx
import os
import glob

def search_kapa_details():
    files = [
        "Supplier_Master_Intelligence_Report.docx",
        "Supplier_Rhythm_Master.docx",
        "OASIS_Supplier_Relational_Audit.docx",
        "Unique_SKU_Ecosystem_Analysis.docx"
    ]
    for fn in files:
        f = os.path.join(r"C:\Users\iLink\.gemini\antigravity\scratch", fn)
        if not os.path.exists(f):
            print(f"File not found: {fn}")
            continue
        try:
            doc = docx.Document(f)
            print(f"\n=== Broad Kapa search in {fn} ===")
            for i, p in enumerate(doc.paragraphs):
                if "KAPA" in p.text.upper():
                    # print paragraph and one surrounding paragraph
                    print(f"P{i}: {p.text}")
            for t_idx, t in enumerate(doc.tables):
                for r_idx, r in enumerate(t.rows):
                    row_text = [c.text.strip() for c in r.cells]
                    row_str = " | ".join(row_text)
                    if "KAPA" in row_str.upper():
                        print(f"Table {t_idx} Row {r_idx}: {row_text}")
        except Exception as e:
            print(f"Error {fn}: {e}")

search_kapa_details()
