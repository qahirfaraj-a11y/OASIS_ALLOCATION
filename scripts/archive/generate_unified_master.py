import pandas as pd
import json
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_unified_master():
    json_path = r'C:\Users\iLink\.gemini\antigravity\scratch\supplier_rhythm_analysis.json'
    fulfillment_xlsx = r'C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Fulfillment_Summary.xlsx'
    intelligence_xlsx = r'C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Intelligence_Report_2025_v3.xlsx'
    output_path = r'C:\Users\iLink\.gemini\antigravity\scratch\Supplier_Master_Intelligence_Report.docx'


    # Load JSON (Order Rhythm)
    with open(json_path, 'r') as f:
        order_rhythm = json.load(f).get('po_rhythm', {})

    # Load Fulfillment (Lead Times)
    df_full = pd.read_excel(fulfillment_xlsx)
    # Clean vendor names for matching (remove SB0009 - etc)
    df_full['MatchName'] = df_full['Vendor Name'].str.split(' - ').str[-1].str.strip()

    # Load Intelligence (Risk/Quality)
    df_intel = pd.read_excel(intelligence_xlsx, sheet_name='Supplier Rankings')

    doc = Document()

    # Title
    title = doc.add_heading('O.A.S.I.S. Supplier Master Intelligence', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This master report reconciles the Ordering Rhythm (frequency of POs) with the Fulfillment Rhythm (speed of delivery) and overall Risk Intelligence.')

    # Key Definitions
    doc.add_heading('Key Definitions', level=1)
    p = doc.add_paragraph()
    p.add_run('Order Rhythm (Median Gap): ').bold = True
    p.add_run('The median number of days between two consecutive Purchase Orders. This determines how often we interact with the supplier.\n')
    p.add_run('Fulfillment Rhythm (Lead Time): ').bold = True
    p.add_run('The median number of days between a PO being issued and the Goods Received Note (GRN). This is extracted from the Fulfillment Summary.\n')
    p.add_run('Total POs vs SKU Count: ').bold = True
    p.add_run('Total POs (Intelligence) counts order events. Fulfillment Count counts individual SKU line items.')

    # Main Table
    doc.add_heading('Supplier Performance Matrix', level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    headers = ['Supplier Name', 'Order Rhythm (Gaps)', 'Fulfillment (Lead Time)', 'Total POs', 'SKU Count', 'Risk Level', 'Quality']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Get union of all suppliers
    all_suppliers = sorted(list(set(order_rhythm.keys()) | set(df_full['MatchName']) | set(df_intel['Supplier Name'])))

    for supplier in all_suppliers:
        # 1. Order Data (from JSON)
        o_data = order_rhythm.get(supplier, {})
        order_gap = o_data.get('median_gap', 'N/A')
        
        # 2. Fulfillment Data
        f_match = df_full[df_full['MatchName'] == supplier]
        lead_time = f_match['median'].iloc[0] if not f_match.empty else 'N/A'
        sku_count = f_match['count'].iloc[0] if not f_match.empty else 'N/A'

        # 3. Intelligence Data
        i_match = df_intel[df_intel['Supplier Name'] == supplier]
        total_pos = i_match['Total POs'].iloc[0] if not i_match.empty else (o_data.get('total_orders', 'N/A'))
        risk = i_match['Risk Level'].iloc[0] if not i_match.empty else 'N/A'
        quality = i_match['Quality'].iloc[0] if not i_match.empty else 'N/A'

        # Filter out if too much missing data (optional, but let's keep all for now)
        if supplier == 'BROOKSIDE DAIRY LIMITED' or not pd.isna(lead_time) or not pd.isna(risk):
            row_cells = table.add_row().cells
            row_cells[0].text = str(supplier)
            row_cells[1].text = f"{order_gap}d" if order_gap != 'N/A' else 'N/A'
            row_cells[2].text = f"{lead_time}d" if lead_time != 'N/A' else 'N/A'
            row_cells[3].text = str(total_pos)
            row_cells[4].text = str(sku_count)
            row_cells[5].text = str(risk)
            row_cells[6].text = str(quality)

    doc.save(output_path)
    print(f"Unified Master saved to {output_path}")

if __name__ == "__main__":
    generate_unified_master()
