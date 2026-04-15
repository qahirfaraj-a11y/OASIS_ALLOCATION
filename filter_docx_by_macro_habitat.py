import os
import sys
import re
import pandas as pd
import zipfile
from docx import Document
from datetime import datetime
from collections import defaultdict

# Paths
SOURCE_DOC = r"C:\Users\iLink\Downloads\OASIS_All_SKUs_Full_Query.docx"
SCORECARD_FILE = r"c:\Users\iLink\.gemini\antigravity\scratch\Full_Product_Allocation_Scorecard_v7.csv"
OUTPUT_DIR = r"c:\Users\iLink\.gemini\antigravity\scratch\oasis\data\filtered_queries"
ANALYSIS_DIR = r"c:\Users\iLink\.gemini\antigravity\scratch\oasis\data\retail_analyses"

def ensure_dirs():
    for d in [OUTPUT_DIR, ANALYSIS_DIR]:
        if not os.path.exists(d):
            os.makedirs(d)

def get_text_from_xml(doc_path):
    """Extremely fast extraction of text from .docx XML with line break preservation."""
    with zipfile.ZipFile(doc_path) as z:
        xml_content = z.read('word/document.xml').decode('utf-8')
        text = xml_content.replace('</w:p>', '\n').replace('<w:br/>', '\n').replace('</w:t>', '\n')
        text = re.sub(r'<.*?>', '', text)
        return text

def extract_metric(text, pattern):
    match = re.search(pattern, text)
    if match:
        return match.group(1).strip()
    return "0"

def filter_by_macro_habitat(target_dept=None, batch_mode=False):
    ensure_dirs()
    print(f"Loading document (OASIS Masterclass 2.0 Parser): {SOURCE_DOC}...")
    full_text = get_text_from_xml(SOURCE_DOC)
    
    # Load Scorecard for name fallback and 100% accuracy
    print(f"Loading Scorecard for cross-referencing: {SCORECARD_FILE}")
    scorecard_df = pd.read_csv(SCORECARD_FILE)
    # Create a mapping of (Revenue, Price) -> Product Name for 100% precision
    # Round metrics to avoid float precision issues during mapping
    scorecard_df['Rounded_Revenue'] = scorecard_df['Total_Revenue'].round(2)
    scorecard_df['Rounded_Price'] = scorecard_df['Unit_Price'].round(2)
    
    # Mapping for fallback: (Rounded_Revenue, Rounded_Price) -> Product Name
    name_lookup = {}
    for _, row in scorecard_df.iterrows():
        key = (row['Rounded_Revenue'], row['Rounded_Price'])
        name_lookup[key] = str(row['Product'])

    # Split by horizontal rule
    chunks = re.split(r'---', full_text)
    print(f"Total chunks found: {len(chunks)}")

    dept_map = defaultdict(list)
    excel_map = defaultdict(list)
    
    # Reconstruct SKU blocks
    i = 1
    while i < len(chunks) - 1:
        yaml_chunk = chunks[i].strip()
        identity_chunk = chunks[i+1].strip()
        
        if "type: SKU" in yaml_chunk:
            habitat_match = re.search(r"department:\s+\"(.*?)\"", yaml_chunk, re.IGNORECASE)
            if not habitat_match:
                habitat_match = re.search(r"Macro-Habitat:\**\s+\[\[(.*?)\]\]", identity_chunk, re.IGNORECASE)
            
            if habitat_match:
                found_dept = habitat_match.group(1).strip().upper()
                full_sku_block = f"---\n{yaml_chunk}\n---\n{identity_chunk}"
                
                # 1. Primary Name Extraction (Identity Chunk)
                extracted_name = "Unknown SKU"
                lines = [l.strip() for l in identity_chunk.split('\n') if l.strip()]
                for l in lines:
                    clean_line = l.replace('#', '').strip()
                    if clean_line and "###" not in l and "---" not in l and "type:" not in l:
                        extracted_name = clean_line
                        break
                
                # 2. Extract Key Metrics for Cross-Ref Logic
                raw_revenue = float(extract_metric(yaml_chunk, r"revenue:\s+([\d\.]+)").replace(',', '') or 0)
                raw_price = float(extract_metric(yaml_chunk, r"price:\s+([\d\.]+)").replace(',', '') or 0)
                
                # 3. Fallback/Verification Case
                if extracted_name == "Unknown SKU" or extracted_name == "":
                    # Attempt to match from Scorecard based on Revenue and Price
                    match_key = (round(raw_revenue, 2), round(raw_price, 2))
                    if match_key in name_lookup:
                        extracted_name = name_lookup[match_key]
                        # print(f"Recovered Name: {extracted_name}")
                
                dept_map[found_dept].append(full_sku_block)
                
                # Full Metric Suite
                qty = extract_metric(yaml_chunk, r"total_quantity:\s+([\d\.]+)")
                ads = extract_metric(yaml_chunk, r"velocity_ads:\s+([\d\.]+)")
                margin = extract_metric(yaml_chunk, r"margin:\s+([\d\-\.]+)")
                supplier = extract_metric(yaml_chunk, r"supplier:\s+\"(.*?)\"")
                
                excel_map[found_dept].append({
                    "SKU Name": extracted_name,
                    "Supplier": supplier,
                    "Price (KES)": raw_price,
                    "Margin (%)": float(margin.replace(',', '') or 0),
                    "Revenue (KES)": raw_revenue,
                    "Total Quantity": float(qty.replace(',', '') or 0),
                    "Velocity (ADS)": float(ads.replace(',', '') or 0)
                })
            i += 2
        else:
            i += 1

    depts_to_process = sorted(dept_map.keys()) if batch_mode else ([target_dept.upper()] if target_dept and target_dept.upper() in dept_map else [])
    
    if not depts_to_process:
        print(f"No SKUs found for target.")
        return

    summary_data = []
    for dept in depts_to_process:
        safe_name = dept.replace(" ", "_").replace("/", "-")
        df = pd.DataFrame(excel_map[dept])
        excel_path = os.path.join(ANALYSIS_DIR, f"OASIS_{safe_name}_Masterclass.xlsx")
        df.to_excel(excel_path, index=False)
        
        summary_data.append({
            "Department": dept, "SKU Count": len(dept_map[dept]),
            "Total Revenue": df["Revenue (KES)"].sum(),
            "Total Quantity": df["Total Quantity"].sum(),
            "Avg Velocity": df["Velocity (ADS)"].mean()
        })

        if not batch_mode or dept in ["BEER", "WINES", "SPIRITS", "CIDERS"]:
            new_doc = Document()
            new_doc.add_heading(f"O.A.S.I.S. Strategic Export: {dept}", 0)
            for block in dept_map[dept][:1000]: # Capacity for Masterclass
                new_doc.add_paragraph(block)
                new_doc.add_paragraph("-" * 20)
            new_doc.save(os.path.join(OUTPUT_DIR, f"OASIS_{safe_name}_Filtered.docx"))

    if batch_mode:
        summary_df = pd.DataFrame(summary_data)
        summary_df.sort_values("Total Revenue", ascending=False).to_excel(os.path.join(ANALYSIS_DIR, "OASIS_RETAIL_SUMMARY_MASTER.xlsx"), index=False)

if __name__ == "__main__":
    target = " ".join(sys.argv[1:]) if len(sys.argv) > 1 else "BEER"
    if target.upper() == "ALL":
        filter_by_macro_habitat(batch_mode=True)
    else:
        filter_by_macro_habitat(target_dept=target)
