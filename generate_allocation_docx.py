import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('OASIS Allocation Engine Breakdown', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    md_path = r"c:\Users\iLink\.gemini\antigravity\scratch\allocation_engine_documentation.md"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('*   ') or line.startswith('-   '):
            p = doc.add_paragraph(line[4:], style='List Bullet')
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('    * ') or line.startswith('    - '):
            p = doc.add_paragraph(line[6:], style='List Bullet 2') # Approximate
        elif line.startswith('|'):
            # Simple table handling (skip complexity for now, add as text)
            p = doc.add_paragraph(line, style='No Spacing')
            p.runs[0].font.name = 'Courier New' # Monospace for tables
        else:
            doc.add_paragraph(line)

    output_path = r"c:\Users\iLink\.gemini\antigravity\scratch\Allocation_Field_Ops_Guide.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
