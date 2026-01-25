import json
import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("CRITICAL: python-docx not installed. Run 'pip install python-docx'")
    exit(1)

# --- CONFIGURATION ---
DATA_DIR = r"C:\Users\iLink\.gemini\antigravity\scratch\app\data"
SALES_FORECAST_JSON = os.path.join(DATA_DIR, "sales_forecasting_2025 (1).json")
OUTPUT_FILE = r"C:\Users\iLink\.gemini\antigravity\scratch\Hayat_Kimya_Logistics_Report_2025_v2.docx"

CATEGORIES = {
    "SANITARY TOWELS": ["SANITARY", "TOWEL", "PAD", "ALWAYS", "KOTEX", "MOLPED", "SOFY", "LIBRESSE"],
    "FABRIC CONDITIONER": ["FABRIC", "CONDITIONER", "SOFTENER", "DOWNY", "STA-SOFT", "COMFORT", "BINGO", "KLEANS"],
    "DIAPERS": ["DIAPER", "PAMPERS", "HUGGIES", "MOLFIX", "SOFTCARE", "BEBEM", "SNOOGGMS", "NAPPY"],
    "WIPES": ["WIPES", "WET", "BABY WIPES"]
}

HAYAT_KEYWORDS = ["HAYAT", "MOLFIX", "MOLPED", "BINGO", "PAPIA", "FAMILIA", "BEBEM"]

# --- DATA HELPERS ---
def normalize(text):
    if not text: return ""
    return str(text).upper().strip()

def classify_category(product_name):
    norm = normalize(product_name)
    for cat, keywords in CATEGORIES.items():
        for k in keywords:
            if k in norm:
                return cat
    return None

def is_hayat(supplier_name, product_name):
    norm_supp = normalize(supplier_name)
    norm_prod = normalize(product_name)
    if "HAYAT" in norm_supp: return True
    if any(k in norm_prod for k in HAYAT_KEYWORDS): return True
    return False

def get_velocity_class(qty):
    if qty >= 50: return "Fast Mover"
    if qty >= 10: return "Medium Mover"
    return "Slow Mover"

# --- DOCX HELPERS ---
def create_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    doc.add_paragraph()

# --- MAIN LOGIC ---
def main():
    print("Loading Data...")
    if not os.path.exists(SALES_FORECAST_JSON):
        print(f"Error: Data file not found at {SALES_FORECAST_JSON}")
        return

    with open(SALES_FORECAST_JSON, 'r') as f:
        sales_forecast = json.load(f)
        
    category_data = {cat: {"hayat": [], "competitors": []} for cat in CATEGORIES}
    
    for product_name, data in sales_forecast.items():
        cat = classify_category(product_name)
        if not cat: continue
        
        qty = data.get('total_10mo_sales', 0)
        velocity = get_velocity_class(qty)
        hayat = is_hayat("Unknown", product_name)
        
        item_data = {"name": product_name, "qty": qty, "velocity": velocity}
        
        if hayat:
            category_data[cat]["hayat"].append(item_data)
        else:
            category_data[cat]["competitors"].append(item_data)

    print("Generating Document...")
    doc = Document()
    
    # Header
    doc.add_heading('Warehousing & Logistics Performance Report', 0)
    p = doc.add_paragraph()
    p.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n').bold = True
    p.add_run('To: Warehousing and Logistics Director\n')
    p.add_run('Subject: Hayat Kimya Product Flow & Competitive Analysis')
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph("This report analyzes the movement of Hayat Kimya products (Molfix, Molped, Bingo, Familia) versus competitors to optimize warehouse slotting and distribution planning.")
    
    total_hayat_vol = sum(sum(x['qty'] for x in v['hayat']) for v in category_data.values())
    doc.add_paragraph(f"Total Hayat Unit Volume (YTD): {total_hayat_vol:,} Units")
    
    # Detailed Analysis
    doc.add_heading('2. Segment Analysis & Warehousing Recommendations', level=1)
    
    # Brand Definitions for detection
    BRAND_MAP = {
        "DIAPERS": {
            "PAMPERS": ["PAMPERS"],
            "HUGGIES": ["HUGGIES"],
            "SOFTCARE": ["SOFTCARE"],
            "R&F / GENERIC": ["R&F", "NAPPY"],
            "MOLFIX (HAYAT)": ["MOLFIX"]
        },
        "SANITARY TOWELS": {
            "ALWAYS": ["ALWAYS"],
            "KOTEX": ["KOTEX"],
            "VELVEX": ["VELVEX"],
            "FAY": ["FAY"],
            "MOLPED (HAYAT)": ["MOLPED"]
        },
        "FABRIC CONDITIONER": {
            "DOWNY": ["DOWNY"],
            "STA-SOFT": ["STA-SOFT", "STASOFT"],
            "COMFORT": ["COMFORT"],
            "BINGO (HAYAT)": ["BINGO"]
        },
        "WIPES": {
            "HANAN": ["HANAN"],
            "VELVEX": ["VELVEX"],
            "FAY": ["FAY"],
            "ARYUV": ["ARYUV"],
            "FAMILIA (HAYAT)": ["FAMILIA"]
        }
    }

    def detect_brand(name, category):
        clean = normalize(name)
        if category in BRAND_MAP:
            for brand, kw_list in BRAND_MAP[category].items():
                if any(k in clean for k in kw_list):
                    return brand
        return "OTHER COMPETITORS"

    for cat, data in category_data.items():
        doc.add_heading(f'Category: {cat}', level=2)
        h_items = sorted(data['hayat'], key=lambda x: x['qty'], reverse=True)
        c_items = sorted(data['competitors'], key=lambda x: x['qty'], reverse=True)
        
        h_vol = sum(x['qty'] for x in h_items)
        c_vol = sum(x['qty'] for x in c_items)
        
        doc.add_paragraph(f"Hayat Volume: {h_vol:,} units | Competitor Volume: {c_vol:,} units")
        
        # Recommendations
        rec = "Low Velocity"
        if h_vol > 500: rec = "High Velocity. Recommend dedicated Pallet Location (Pick Face)."
        elif h_vol > 100: rec = "Medium Velocity. Recommend Case Flow or Shelving."
        else: rec = "Low Velocity. Recommend Upper Steel / Reserve."
        doc.add_paragraph(f"Logistics Recommendation: {rec}").runs[0].bold = True
        
        # --- COMPETITOR BREAKDOWN BY BRAND ---
        doc.add_heading("Competitor Breakdown by Brand", level=3)
        
        # Group competitors by brand
        comp_brands = {} # Brand -> list of items
        for item in c_items:
            brand = detect_brand(item['name'], cat)
            if brand not in comp_brands: comp_brands[brand] = []
            comp_brands[brand].append(item)
            
        # Iterate through major brands and show gap
        for brand, items in sorted(comp_brands.items(), key=lambda x: sum(i['qty'] for i in x[1]), reverse=True):
            brand_vol = sum(i['qty'] for i in items)
            doc.add_paragraph(f"VS {brand} (Vol: {brand_vol:,})", style='Heading 4')
            
            # Create comparison table
            # Columns: Competitor Item | Qty | Closest Hayat Item (by Rank) | Qty | Gap
            table_rows = []
            
            # Show top 5 items for this brand
            limit = min(len(items), 10)
            for i in range(limit):
                c_itm = items[i]
                # Compare to Hayat equivalent rank if available, else just top Hayat
                h_itm = h_items[i] if i < len(h_items) else {"name": "-", "qty": 0}
                
                gap = c_itm['qty'] - h_itm['qty']
                gap_str = f"-{gap:,}" if gap > 0 else f"+{abs(gap):,}"
                
                table_rows.append([c_itm['name'], c_itm['qty'], h_itm['name'], h_itm['qty'], gap_str])
                
            if table_rows:
                create_table(doc, [f"{brand} Item", "Vol", "Hayat Item", "Vol", "Gap"], table_rows)
            doc.add_paragraph() # Spacer

    # Strategic Implication
    doc.add_heading('3. Strategic Implications', level=1)
    doc.add_paragraph("DIAPERS (MOLFIX): Primary logistics driver. Ensure Fast Movers (Pants) are at ergonomic height.")
    doc.add_paragraph("SANITARY (MOLPED): Moderate volume. Co-locate with Diapers.")
    doc.add_paragraph("CHEMICALS (BINGO) & WIPES: Low volume. Consolidate in slow-moving aisles.")

    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
