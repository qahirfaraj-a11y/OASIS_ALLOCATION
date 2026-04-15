from docx import Document
import pandas as pd

doc_path = r"C:\Users\iLink\Downloads\OASIS_All_SKUs_Full_Query.docx"
doc = Document(doc_path)

# Try to find tables first
if doc.tables:
    print(f"Found {len(doc.tables)} tables.")
    # Peek at the first few rows of the first table
    table = doc.tables[0]
    for row in table.rows[:5]:
        print([cell.text for cell in row.cells])
else:
    print("No tables found. Peeking at paragraphs...")
    for para in doc.paragraphs[:20]:
        if para.text.strip():
            print(para.text)
