import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_currency(value):
    return f"KES {value:,.2f}"

def generate_masterclass_2_0_report():
    print("Generating Operational Masterclass 2.0 Report...")
    
    with open('alcohol_masterclass_2_0_intel.json', 'r') as f:
        data = json.load(f)
        
    RETAIL_DIR = r"c:\Users\iLink\.gemini\antigravity\scratch\oasis\data\retail_analyses"
    
    doc = Document()
    
    # 1. Executive Branding
    title = doc.add_heading('OPERATIONAL MASTERCLASS 2.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Strategic Alcohol Portfolio & Neural Integration', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Neural Suite 2026 | {data['timestamp']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" + "="*70 + "\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Portfolio Strategy
    doc.add_heading('1. The Alcohol Portfolio: Structural Roles', level=1)
    doc.add_paragraph(
        "This masterclass defines the specific financial and operational roles of each sector. "
        "A retail section is not just a collection of products; it is a system of capital flow."
    )
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    h = table.rows[0].cells
    h[0].text = 'Sector'
    h[1].text = 'Role'
    h[2].text = 'SKU Count'
    h[3].text = 'Total Revenue (KES)'
    
    sectors = [
        ('BEER', 'Velocity Anchor', data['sector_analyses']['BEER']['sku_count'], format_currency(data['sector_analyses']['BEER']['rev_total'])),
        ('WINES', 'Complexity Risk', data['sector_analyses']['WINES']['sku_count'], format_currency(data['sector_analyses']['WINES']['rev_total'])),
        ('SPIRITS', 'Margin/Security', data['sector_analyses']['SPIRITS']['sku_count'], format_currency(data['sector_analyses']['SPIRITS']['rev_total'])),
    ]
    
    for sect, role, count, rev in sectors:
        row = table.add_row().cells
        row[0].text = sect
        row[1].text = role
        row[2].text = str(count)
        row[3].text = rev

    # 3. Beer Section: The Velocity Shield
    doc.add_heading('2. Beer: The Defensive Wall', level=1)
    doc.add_paragraph(
        "Beer is your highest-velocity sector. These items drive 70%+ of your alcohol revenue. "
        "The following 'Group A' SKUs require a zero-stock-out mandate."
    )
    
    for item in data['sector_analyses']['BEER']['velocity_leaders'][:5]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{item['SKU Name']} ").bold = True
        p.add_run(f"| ADS: {item['Velocity (ADS)']:.1f} units | Rev: {format_currency(item['Revenue (KES)'])}")
        
    # 4. Wine Section: The Capital Audit
    doc.add_heading('3. Wine: The Capital Efficiency Audit', level=1)
    doc.add_paragraph(
        f"The Wine section carries {data['sector_analyses']['WINES']['sku_count']} SKUs. "
        f"Critically, **{data['sector_analyses']['WINES']['c_class_count']} SKUs** (The Long Tail) "
        f"contribute only **{data['sector_analyses']['WINES']['c_class_rev_pct']:.1f}%** of total wine revenue."
    )
    doc.add_paragraph("ACTION: Execute a 15% SKU rationalization on Group C items to free up capital and shelf space.")
    
    # 5. Neural Substitution Shield
    doc.add_heading('4. Neural Substitution Shield (Ghost Demand Prevention)', level=1)
    doc.add_paragraph(
        "If a customer cannot find their primary brand, they will either switch to a 'Neural Substitute' or leave the store. "
        "This map defines the closest mathematical alternatives for your top sellers."
    )
    
    for sku, info in data['substitution_shield'].items():
        doc.add_heading(f"Anchor: {sku}", level=2)
        doc.add_paragraph(f"Primary Substitutes: {', '.join(info['substitutes'])}")
        doc.add_paragraph(f"Cross-Category Halo: {', '.join(info['halo'])}")
        
    # 6. Supplier Resilience
    doc.add_heading('5. Supplier Resilience Appendix', level=1)
    doc.add_paragraph("Auditing the reliability and fulfillment stability of your primary alcohol partners.")
    
    table_s = doc.add_table(rows=1, cols=3)
    table_s.style = 'Light List Accent 1'
    hs = table_s.rows[0].cells
    hs[0].text = 'Supplier'
    hs[1].text = 'Reliability (%)'
    hs[2].text = 'Lead Time Var (Days)'
    
    for name, info in data['supplier_audit'].items():
        rs = table_s.add_row().cells
        rs[0].text = name
        rs[1].text = f"{info['Supplier_Reliability'] * 100:.1f}%"
        rs[2].text = f"{info['Lead_Time_Days']:.1f} days avg"

    report_path = os.path.join(RETAIL_DIR, "Alcohol_Operational_Masterclass_2_0.docx")
    doc.save(report_path)
    print(f"Masterclass 2.0 Finalized: {report_path}")

if __name__ == "__main__":
    generate_masterclass_2_0_report()
