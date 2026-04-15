"""
Generates the O.A.S.I.S. Client Implementation Playbook as a styled Word document.
Run once to produce the .docx file.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _styled_para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def build_playbook():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Cover
    doc.add_paragraph()
    doc.add_paragraph()
    _styled_para(doc, 'O.A.S.I.S.', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _styled_para(doc, 'Client Implementation Playbook', bold=True, size=20,
                 color=RGBColor(0x2C, 0x3E, 0x50), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _styled_para(doc, 'From First Contact to Full Autonomous Replenishment', italic=True, size=14,
                 color=RGBColor(0x7F, 0x8C, 0x8D), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    _styled_para(doc, 'INTERNAL - iLink Operations', bold=True, size=12,
                 color=RGBColor(0xCC, 0x00, 0x00), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # TOC
    doc.add_heading('Table of Contents', level=1)
    toc = [
        'Phase 0: Prospecting & First Contact',
        'Phase 1: The Forensic Audit (The Free Diagnosis)',
        'Phase 2: API Hook & Shadow Mode',
        'Phase 3: The AMIT Flush (Dead Stock Liquidation)',
        'Phase 4: High-Velocity Hyper-Funding (DHARAM)',
        'Phase 5: Supplier Shield Activation (LATA)',
        'Phase 6: Full Autonomous Ordering (MANDE)',
        'Post-Implementation: Ongoing Value',
        'Appendix: Data Requirements Reference',
    ]
    for item in toc:
        _styled_para(doc, item, size=11, space_after=3)
    doc.add_page_break()

    # ---- PHASE 0 ----
    doc.add_heading('Phase 0: Prospecting & First Contact', level=1)
    _styled_para(doc, 'Timeline: Day 0', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'Identify a retail operation likely hemorrhaging working capital due to manual procurement, '
        'and secure agreement to receive a free forensic audit.'
    )
    doc.add_heading('The Hook', level=2)
    doc.add_paragraph(
        '"We would like to run a free, no-obligation operations audit on your store. We will take your raw '
        'sales and purchasing data, run it through our forensic engine, and show you exactly how much money '
        'your operation is losing to dead stock, supplier inconsistency, and logistics friction. The audit '
        'takes 24 hours. You lose nothing. You keep the report regardless."',
        style='Intense Quote'
    )
    doc.add_heading('Data Request', level=2)
    doc.add_paragraph('Upon agreement, formally request the following raw data exports:')
    _add_table(doc,
        ['#', 'Data Category', 'Required Fields', 'Purpose'],
        [
            ['1', 'POS / Sales Log', 'Item Name, Qty Sold, Date, Price (optional)', 'Daily velocity & dead stock detection'],
            ['2', 'GRN / Inbound Log', 'Supplier, Item, PO Qty, GRN Qty, Date', 'Supplier fulfillment & lead time variance'],
            ['3', 'Returns (PRTS/GRTS)', 'Item, Qty Adjusted, Reason, Net Amount', 'Spoilage & wastage quantification'],
            ['4', 'Branch Transfers', 'From Branch, To Branch, Item, Qty, Cost', 'Allocation failure measurement'],
            ['5', 'Stock Snapshot (Optional)', 'Item, Current Qty, Cost Price', 'Current stockout/overstock identification'],
        ]
    )
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Data received. Proceed once at least POS and GRN logs are in hand.', bold=True)
    doc.add_page_break()

    # ---- PHASE 1 ----
    doc.add_heading('Phase 1: The Forensic Audit (The Free Diagnosis)', level=1)
    _styled_para(doc, 'Timeline: 24-48 Hours After Data Receipt', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'Ingest the prospect\'s raw data, execute the full O.A.S.I.S. forensic engine, and produce a '
        'consulting-grade report quantifying total revenue bleed.'
    )
    doc.add_heading('Engine Execution Methodology', level=2)
    _add_table(doc,
        ['Engine', 'Detection Rule', 'What It Finds', 'Metric Produced'],
        [
            ['AMIT', 'ADS < 0.2 units/day AND SOH > 15 units', 'Dead stock trapping working capital', 'KES trapped per item'],
            ['DHARAM', 'ADS > 2.0 units/day AND SOH = 0', 'Fast movers out of stock', 'KES lost revenue (14-day projection)'],
            ['LATA', 'Fulfillment < 85% OR Lead Var > 3 days SD', 'Hostile/unreliable suppliers', 'Fulfillment % and variance per vendor'],
            ['MANDE', 'Sum of Net Amount from PRTS + STI docs', 'Returns and transfer costs', 'Total KES network entropy'],
        ]
    )
    doc.add_heading('Deliverables', level=2)
    doc.add_paragraph('Executive Summary Word Document (OASIS_Executive_Diagnostic.docx)', style='List Bullet')
    doc.add_paragraph('Raw Forensic Data Excel (OASIS_Forensic_Audit_Data.xlsx)', style='List Bullet')
    doc.add_paragraph('Live Streamlit dashboard presentation (45-min meeting)', style='List Bullet')
    doc.add_heading('The Pitch Close', level=2)
    doc.add_paragraph(
        '"This is what your business looks like under a microscope. These are not opinions. These are mathematical '
        'facts extracted from your own data. The question is: do you want to keep bleeding, or do you want us to fix it?"',
        style='Intense Quote'
    )
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Client signs contract. Proceed to Shadow Mode.', bold=True)
    doc.add_page_break()

    # ---- PHASE 2 ----
    doc.add_heading('Phase 2: API Hook & Shadow Mode', level=1)
    _styled_para(doc, 'Timeline: Weeks 1-2 After Contract', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'Connect O.A.S.I.S. to the client\'s live ERP/POS system. Run in Shadow Mode: generate daily POs '
        'internally without sending them to suppliers, while the human buyers continue ordering normally.'
    )
    doc.add_heading('Process', level=2)
    doc.add_paragraph('Establish automated daily data pull from client ERP (SQL, API, or SFTP).', style='List Number')
    doc.add_paragraph('O.A.S.I.S. generates daily Shadow POs. These are NOT sent to suppliers.', style='List Number')
    doc.add_paragraph('Generate Daily Shadow Comparison: Human buyer orders vs. O.A.S.I.S. orders.', style='List Number')
    doc.add_paragraph('After 14 days, present aggregated Shadow Review to client executives.', style='List Number')
    doc.add_heading('The Shadow Review Presentation', level=2)
    doc.add_paragraph(
        'Show the client the divergences: items the buyer over-ordered (future dead stock), items the buyer '
        'missed (future stockouts), hostile suppliers the buyer used without adjusting quantities. This builds '
        'absolute trust in the algorithm before it takes control.'
    )
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Client approves transition from Shadow to Active Mode.', bold=True)
    doc.add_page_break()

    # ---- PHASE 3 ----
    doc.add_heading('Phase 3: The AMIT Flush (Dead Stock Liquidation)', level=1)
    _styled_para(doc, 'Timeline: Week 3', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'Before O.A.S.I.S. starts buying new stock, stop the bleeding. Activate the AMIT engine to '
        'systematically liquidate trapped capital in dormant inventory.'
    )
    doc.add_heading('Process', level=2)
    doc.add_paragraph('Generate the AMIT Negative List (all dead SKUs). Present for client sign-off.', style='List Number')
    doc.add_paragraph('Activate system-level purchase blocks on Negative List items.', style='List Number')
    doc.add_paragraph('Execute liquidation: promotional pricing, formal write-offs, or supplier returns (PRTS).', style='List Number')
    doc.add_paragraph('Track and report weekly Capital Recovery to client finance team.', style='List Number')
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Capital begins freeing up. Client authorizes active purchasing for fast movers.', bold=True)
    doc.add_page_break()

    # ---- PHASE 4 ----
    doc.add_heading('Phase 4: High-Velocity Hyper-Funding (DHARAM)', level=1)
    _styled_para(doc, 'Timeline: Weeks 4-6', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'Redirect freed capital into hyper-funding the top 20% fastest-moving SKUs (the Revenue Core). '
        'Ensure zero stockouts on items generating 80% of revenue.'
    )
    doc.add_heading('Order Quantity Formula', level=2)
    doc.add_paragraph('Order Qty = (ADS x Lead Time Days) + Safety Stock - Current SOH - In-Transit Qty')
    doc.add_paragraph(
        'Safety Stock is dynamically calculated per item based on standard deviation of daily sales. '
        'POs are grouped by supplier and presented for daily one-click approval.'
    )
    doc.add_heading('Success Metric', level=2)
    _styled_para(doc, 'Target: 0% stockout rate on Revenue Core items within 14 days of activation.', bold=True)
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Revenue Core stockout rate hits 0%. Client authorizes full catalog expansion.', bold=True)
    doc.add_page_break()

    # ---- PHASE 5 ----
    doc.add_heading('Phase 5: Supplier Shield Activation (LATA)', level=1)
    _styled_para(doc, 'Timeline: Month 2', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'With the floor optimized, fix the supply chain. Activate the LATA engine to dynamically manage '
        'safety stock buffers per supplier based on their measured reliability.'
    )
    doc.add_heading('Safety Stock Multipliers', level=2)
    _add_table(doc,
        ['Supplier Classification', 'Safety Stock Multiplier', 'Logic'],
        [
            ['RELIABLE (Green)', '1.2x', 'Lean ordering. Minimal buffer required.'],
            ['WATCH (Yellow)', '1.5x', 'Moderate buffer for inconsistent delivery.'],
            ['HOSTILE (Red)', '2.0x+', 'Aggressive buffer to protect shelf from short-ships.'],
        ]
    )
    doc.add_heading('Decision Gate', level=2)
    _styled_para(doc, 'Safety stock dynamically managed. Shelf availability improves without capital inflation.', bold=True)
    doc.add_page_break()

    # ---- PHASE 6 ----
    doc.add_heading('Phase 6: Full Autonomous Ordering (MANDE)', level=1)
    _styled_para(doc, 'Timeline: Month 3', italic=True, color=RGBColor(0x7F, 0x8C, 0x8D))
    doc.add_heading('Objective', level=2)
    doc.add_paragraph(
        'O.A.S.I.S. takes over 100% of the daily ordering cycle across all branches and all SKUs. '
        'The MANDE engine mathematically allocates inbound stock to each branch based on localized demand.'
    )
    doc.add_heading('Procurement Team Transition', level=2)
    doc.add_paragraph(
        'The buying team transitions from manual order creation to order approval. Their new workflow: '
        'Open dashboard, review auto-generated Daily Master PO, click Approve. Exception management only.'
    )
    doc.add_heading('Continuous Optimization', level=2)
    doc.add_paragraph('All engine parameters re-calculate daily based on latest POS and GRN data.', style='List Bullet')
    doc.add_paragraph('AMIT scan runs weekly, auto-adding newly dormant items to the Negative List.', style='List Bullet')
    doc.add_paragraph('LATA scan runs daily, adjusting supplier risk scores as new GRNs arrive.', style='List Bullet')
    doc.add_paragraph('Demand model improves with each cycle (15-25% more accurate after 90 days).', style='List Bullet')
    doc.add_heading('Final Outcome', level=2)
    _styled_para(doc, 'Lateral transfers approach zero. Stockouts < 2%. Dead stock < 5%. The operation is self-driving.', bold=True)
    doc.add_page_break()

    # ---- POST-IMPLEMENTATION ----
    doc.add_heading('Post-Implementation: Ongoing Value', level=1)
    _add_table(doc,
        ['Metric', 'Pre-O.A.S.I.S.', 'Post-O.A.S.I.S. Target'],
        [
            ['Dead Stock %', '30-50% of catalog', '< 5%'],
            ['Stockout Rate (Fast Movers)', '10-20%', '< 2%'],
            ['Capital Utilization', '50-70%', '> 95%'],
            ['Supplier Fulfillment (Avg)', 'Unmonitored', '> 85% enforced'],
            ['Lateral Transfers', 'High (weekly)', 'Near zero'],
            ['Procurement Team', '5-10 manual buyers', '1-2 approval managers'],
            ['Order Accuracy', 'Human gut instinct', 'Mathematical precision'],
        ]
    )
    doc.add_page_break()

    # ---- APPENDIX ----
    doc.add_heading('Appendix: Timeline Summary', level=1)
    _add_table(doc,
        ['Week', 'Phase', 'Key Deliverable', 'Decision Gate'],
        [
            ['Day 0', 'First Contact', 'Data Request', 'Data received'],
            ['Day 1-2', 'Phase 1: Forensic Audit', 'Word Report + Excel', 'Contract signed'],
            ['Week 1-2', 'Phase 2: Shadow Mode', '14-day comparison', 'Active mode approved'],
            ['Week 3', 'Phase 3: AMIT Flush', 'Capital Recovery Reports', 'Active purchasing authorized'],
            ['Week 4-6', 'Phase 4: DHARAM', 'Revenue Impact Reports', 'Full catalog authorized'],
            ['Month 2', 'Phase 5: LATA', 'Supplier Scorecards', 'Network optimization approved'],
            ['Month 3', 'Phase 6: Full Autonomous', 'Zero-touch procurement', 'Self-driving operation'],
        ]
    )

    _styled_para(doc, 'END OF PLAYBOOK', bold=True, size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x7F, 0x8C, 0x8D))

    return doc


if __name__ == '__main__':
    doc = build_playbook()
    output_path = 'OASIS_Client_Implementation_Playbook.docx'
    doc.save(output_path)
    print(f'Playbook saved to {output_path}')
