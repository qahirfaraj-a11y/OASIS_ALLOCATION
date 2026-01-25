import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file_path, docx_file_path):
    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: Could not find file {md_file_path}")
        return

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=2)
        
        # Bold/Strong emphasis (simple check)
        elif '**' in line:
            p = document.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Bullet points
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        
        # Tables (Very basic handling - just text for now to avoid complexity)
        elif '|' in line:
             # Skip separator lines like |---|---|
            if set(line.replace('|', '').replace('-', '').replace(' ', '').replace(':', '')) == set():
                continue
            document.add_paragraph(line, style='No Spacing') # Use monospaced look or tight spacing
            
        else:
            document.add_paragraph(line)

    document.save(docx_file_path)
    print(f"Successfully created: {docx_file_path}")

if __name__ == "__main__":
    md_path = r"C:\Users\iLink\.gemini\antigravity\brain\727303c1-e050-46d0-8394-7f5849b8a103\maccuisine_intelligence_report.md"
    docx_path = r"C:\Users\iLink\.gemini\antigravity\scratch\Maccuisine_Intelligence_Analysis.docx"
    markdown_to_docx(md_path, docx_path)
