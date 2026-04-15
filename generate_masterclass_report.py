import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_currency(value):
    return f"KES {value:,.2f}"

def generate_masterclass_report():
    print("Generating Operational Masterclass Report...")
    
    with open('alcohol_masterclass_intel.json', 'r') as f:
        data = json.load(f)
        
    ANALYSIS_DIR = r"c:\Users\iLink\.gemini\antigravity\scratch\oasis\data\retail_analyses"
    
    doc = Document()
    
    # 1. Professional Title Page
    title = doc.add_heading('OPERATIONAL MASTERCLASS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Alcohol Section Strategy & Inventory Optimization', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n" + "="*60 + "\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Section Ecosystem Overview
    doc.add_heading('1. The Alcohol Ecosystem: Sector Roles', level=1)
    doc.add_paragraph(
        "This breakdown defines the strategic role of each sub-category within your store's neural network. "
        "Understanding these roles is the first step toward operational mastery."
    )
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sector'
    hdr_cells[1].text = 'SKU Count'
    hdr_cells[2].text = 'Revenue Contribution'
    hdr_cells[3].text = 'Operational Role'
    
    rows = [
        ('BEER', data['summary']['total_beers'], format_currency(data['summary']['top_beer_revenue']), 'Traffic Anchor / Cash-Flow Driver'),
        ('WINES', data['summary']['total_wines'], 'KES 2.87M', 'Premium Variety / Capacity Risk'),
        ('SPIRITS', data['summary']['total_spirits'], 'KES 0.66M', 'Security Focus / High Value-to-Volume'),
    ]
    
    for sect, count, rev, role in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = sect
        row_cells[1].text = str(count)
        row_cells[2].text = rev
        row_cells[3].text = role

    # 3. Beer Section: The High-Velocity Shield
    doc.add_heading('2. The Beer Section: 100% Availability Mandate', level=1)
    doc.add_paragraph(
        "Beer represents your primary footway driver. Any stock-out here creates 'Ghost Demand' loss for other categories. "
        "The following SKUs are your 'Velocity Shield'—they must never be empty."
    )
    
    table_b = doc.add_table(rows=1, cols=4)
    table_b.style = 'Light Shading Accent 1'
    h = table_b.rows[0].cells
    h[0].text = 'Anchor SKU'
    h[1].text = 'ADS (Daily Sales)'
    h[2].text = 'Total Revenue (KES)'
    h[3].text = 'Refill Frequency'
    
    for item in data['beer_pareto'][:8]:
        rc = table_b.add_row().cells
        rc[0].text = item['SKU Name']
        rc[1].text = f"{item['Velocity (ADS)']:.1f} units"
        rc[2].text = format_currency(item['Revenue (KES)'])
        # High velocity refill logic
        rc[3].text = "24-48 Hours" if item['Velocity (ADS)'] > 20 else "Weekly"
        
    doc.add_paragraph(
        "\nSTRATEGIC INSIGHT: Identify the peak sales window (likely Friday AM) and ensure shelf-stock is doubled before this window starts."
    )

    # 4. Wine Section: The Complexity Risk
    doc.add_heading('3. The Wine Section: Capital Efficiency & SKU Bloat', level=1)
    doc.add_paragraph(
        f"The Wine section carries {data['summary']['total_wines']} unique labels. "
        f"Critically, **{data['summary']['wine_dead_pct']}%** of these SKUs ({data['summary']['wine_dead_stock']} labels) "
        "have a daily velocity of less than 0.05 units per day. This is 'Dead Stock'."
    )
    
    doc.add_heading('The High-Performing Exceptions (Wines)', level=2)
    for item in data['wine_pareto'][:5]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{item['SKU Name']} ").bold = True
        p.add_run(f"({format_currency(item['Revenue (KES)'])})")

    doc.add_paragraph(
        "OPERATIONAL GOAL: Begin a phased 'Catalog Clean-up'. Delist the bottom 100 labels that have not moved in 90 days "
        "and consolidate that shelf space for more units of your top-performing beers or premium spirits."
    )

    # 5. Neural Halo Placement
    doc.add_heading('4. Neural Network Placement (Halo Mapping)', level=1)
    doc.add_paragraph(
        "Our neural graph tracks 'Affinity Bridges'—items purchased in the same basket as your Alcohol Anchors. "
        "Placing these items within 10 meters of the alcohol section increases basket conversion."
    )
    
    doc.add_heading('The Alcohol Buddy List', level=2)
    for target, weight in data['halo_partners'].items():
        if target not in ["Retail Market", "UNKNOWN"]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(target)
    
    doc.add_paragraph(
        "ACTION: Place high-margin snacks (chips, nuts) and premium mixers (tonic, sodas) directly adjacent to the Beer chiller."
    )

    # 6. Safety Stock Formula
    doc.add_heading('5. Scientific Order Triggers', level=1)
    doc.add_paragraph(
        "To achieve operational excellence, order quantities should follow the O.A.S.I.S. formula:\n"
        "Safety Stock = (Max Lead Time - Avg Lead Time) * Average Daily Sales"
    )
    
    doc.add_paragraph("\n" + "="*60)
    footer = doc.add_paragraph("END OF STRATEGIC OVERVIEW")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_path = os.path.join(ANALYSIS_DIR, "Alcohol_Operational_Masterclass.docx")
    doc.save(report_path)
    print(f"Masterclass Report Generated: {report_path}")

if __name__ == "__main__":
    generate_masterclass_report()
