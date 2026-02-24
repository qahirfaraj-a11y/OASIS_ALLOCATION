"""
Convert Markdown Value Proposition to Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Read the markdown file
md_path = r"C:\Users\iLink\.gemini\antigravity\brain\4fcb6031-7a41-4f53-bb16-138c775d8081\chandarana_value_proposition.md"
output_path = r"C:\Users\iLink\.gemini\antigravity\scratch\Chandarana_OASIS_Value_Proposition.docx"

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper function to add formatted paragraph
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_from_text(doc, table_text):
    """Parse markdown table and add to document"""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip() and not l.strip().startswith('|--')]
    if not lines:
        return
    
    # Parse header
    header = [c.strip() for c in lines[0].split('|') if c.strip()]
    rows = []
    for line in lines[1:]:
        if '---' in line:
            continue
        row = [c.strip() for c in line.split('|') if c.strip()]
        if row:
            rows.append(row)
    
    if not header:
        return
        
    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        if i < len(hdr_cells):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell in enumerate(row):
            if j < len(row_cells):
                row_cells[j].text = cell
    
    doc.add_paragraph()  # Add space after table

# Process content
lines = content.split('\n')
i = 0
in_code_block = False
code_content = []

while i < len(lines):
    line = lines[i]
    
    # Code block handling
    if line.strip().startswith('```'):
        if in_code_block:
            # End code block - add as formatted text
            code_text = '\n'.join(code_content)
            p = doc.add_paragraph()
            p.style = 'No Spacing'
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            doc.add_paragraph()
            code_content = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_content.append(line)
        i += 1
        continue
    
    # Headers
    if line.startswith('# '):
        add_heading(doc, line[2:].strip(), 1)
    elif line.startswith('## '):
        add_heading(doc, line[3:].strip(), 2)
    elif line.startswith('### '):
        add_heading(doc, line[4:].strip(), 3)
    
    # Blockquotes
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.style = 'Intense Quote'
        p.add_run(line[2:].strip())
    
    # Tables
    elif line.strip().startswith('|') and '|' in line:
        # Collect all table lines
        table_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip().startswith('|'):
            table_lines.append(lines[j])
            j += 1
        add_table_from_text(doc, '\n'.join(table_lines))
        i = j - 1  # Skip processed lines
    
    # Bullet points
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')
    
    # Numbered lists
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')
    
    # Horizontal rules
    elif line.strip() == '---':
        doc.add_paragraph('_' * 50)
    
    # Regular paragraphs
    elif line.strip():
        # Clean up markdown formatting
        text = line.strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.+?)`', r'\1', text)  # Code
        p = doc.add_paragraph(text)
    
    i += 1

# Save document
doc.save(output_path)
print(f"Word document saved to: {output_path}")
